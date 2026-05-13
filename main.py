from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from kivy.uix.gridlayout import GridLayout
from kivy.uix.scrollview import ScrollView
from kivy.graphics import Color, RoundedRectangle
from kivy.clock import Clock

from pathlib import Path
from datetime import datetime
import json
import webbrowser
from urllib.parse import quote
import re
import threading
import queue
import tempfile
import os

import sounddevice as sd
from scipy.io.wavfile import write
import speech_recognition as sr
import numpy as np
import soundfile as sf

APP_NAME = "Mi Agenda Android"
BASE_DIR = Path.home() / "MiAgendaAndroid"
ENTRIES_DIR = BASE_DIR / "entradas"
EXPORTS_DIR = BASE_DIR / "exportados"
AUDIO_DIR = BASE_DIR / "audios"
CONFIG_FILE = BASE_DIR / "config.json"

ENTRIES_DIR.mkdir(parents=True, exist_ok=True)
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
AUDIO_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_CONFIG = {
    "idioma_transcripcion": "es-ES",
    "correccion_automatica": True,
    "punto_final": True,
    "mayusculas": True,
    "carpeta_exportacion": str(EXPORTS_DIR)
}

if not CONFIG_FILE.exists():
    CONFIG_FILE.write_text(
        json.dumps(DEFAULT_CONFIG, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )

try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    PDF_OK = True
except Exception:
    PDF_OK = False


def load_config():
    try:
        data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        cfg = DEFAULT_CONFIG.copy()
        cfg.update(data)
        return cfg
    except Exception:
        return DEFAULT_CONFIG.copy()


def save_config(data):
    CONFIG_FILE.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )


def style_button(btn, bg=(0.20, 0.45, 0.90, 1), color=(1, 1, 1, 1)):
    btn.background_normal = ""
    btn.background_down = ""
    btn.background_color = bg
    btn.color = color


class Card(BoxLayout):
    def __init__(self, bg_color=(0.12, 0.12, 0.16, 1), radius=16, **kwargs):
        super().__init__(**kwargs)
        with self.canvas.before:
            Color(*bg_color)
            self.rect = RoundedRectangle(pos=self.pos, size=self.size, radius=[radius])
        self.bind(pos=self._update_rect, size=self._update_rect)

    def _update_rect(self, *args):
        self.rect.pos = self.pos
        self.rect.size = self.size


class SelectListPopup(Popup):
    def __init__(self, title_text, items, on_select, **kwargs):
        super().__init__(title=title_text, size_hint=(0.92, 0.88), **kwargs)
        self.on_select_callback = on_select

        root = BoxLayout(orientation="vertical", spacing=8, padding=10)

        if not items:
            root.add_widget(Label(text="No hay elementos disponibles."))
            btn = Button(text="Cerrar", size_hint_y=None, height=45)
            style_button(btn, (0.22, 0.44, 0.82, 1))
            btn.bind(on_press=self.dismiss)
            root.add_widget(btn)
            self.content = root
            return

        scroll = ScrollView()
        content = GridLayout(cols=1, spacing=6, size_hint_y=None)
        content.bind(minimum_height=content.setter("height"))

        for idx, item in enumerate(items):
            btn = Button(
                text=f"{idx + 1}. {item.name}",
                size_hint_y=None,
                height=48
            )
            style_button(btn, (0.18, 0.20, 0.26, 1))
            btn.bind(on_press=lambda instance, i=idx: self.select_item(i))
            content.add_widget(btn)

        scroll.add_widget(content)
        root.add_widget(scroll)

        btn_close = Button(text="Cerrar", size_hint_y=None, height=45)
        style_button(btn_close, (0.78, 0.25, 0.25, 1))
        btn_close.bind(on_press=self.dismiss)
        root.add_widget(btn_close)

        self.content = root

    def select_item(self, index):
        self.on_select_callback(index)
        self.dismiss()


class AgendaLayout(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(orientation="horizontal", spacing=12, padding=12, **kwargs)

        self.config_data = load_config()
        self.ultimo_audio = None
        self.grabando = False
        self.fs = 16000
        self.audio_chunks = []
        self.input_stream = None
        self.historial_notas = []
        self.historial_audios = []
        self.archivo_actual = None
        self.reproduciendo = False

        # Transcripción en vivo
        self.transcribiendo_en_vivo = False
        self.live_audio_buffer = []
        self.live_chunk_seconds = 5
        self.live_queue = queue.Queue()
        self.live_worker_thread = None
        self.live_transcript_items = []
        self.live_last_text = ""

        with self.canvas.before:
            Color(0.07, 0.08, 0.10, 1)
            self.bg = RoundedRectangle(pos=self.pos, size=self.size, radius=[0])
        self.bind(pos=self._update_bg, size=self._update_bg)

        izquierda = Card(
            orientation="vertical",
            size_hint_x=0.30,
            spacing=8,
            padding=10,
            bg_color=(0.11, 0.12, 0.16, 1)
        )

        izquierda.add_widget(Label(
            text="Historial de notas",
            size_hint_y=None,
            height=34,
            color=(1, 1, 1, 1),
            bold=True
        ))

        self.lista_notas = TextInput(
            readonly=True,
            multiline=True,
            background_color=(0.08, 0.09, 0.12, 1),
            foreground_color=(1, 1, 1, 1)
        )
        izquierda.add_widget(self.lista_notas)

        for texto, color, accion in [
            ("Actualizar historial", (0.22, 0.44, 0.82, 1), self.actualizar_historial_notas),
            ("Abrir nota", (0.18, 0.62, 0.42, 1), self.abrir_nota_desde_lista),
            ("Borrar nota", (0.78, 0.25, 0.25, 1), self.borrar_nota_desde_lista),
            ("Editar título", (0.62, 0.42, 0.15, 1), self.editar_titulo_desde_lista),
            ("Renombrar nota", (0.52, 0.40, 0.82, 1), self.renombrar_nota_desde_lista),
        ]:
            btn = Button(text=texto, size_hint_y=None, height=42)
            style_button(btn, color)
            btn.bind(on_press=accion)
            izquierda.add_widget(btn)

        izquierda.add_widget(Label(
            text="Audios grabados",
            size_hint_y=None,
            height=34,
            color=(1, 1, 1, 1),
            bold=True
        ))

        self.lista_audios = TextInput(
            readonly=True,
            multiline=True,
            background_color=(0.08, 0.09, 0.12, 1),
            foreground_color=(1, 1, 1, 1)
        )
        izquierda.add_widget(self.lista_audios)

        for texto, color, accion in [
            ("Actualizar audios", (0.22, 0.44, 0.82, 1), self.actualizar_historial_audios),
            ("Reproducir audio", (0.12, 0.70, 0.32, 1), self.reproducir_audio_desde_lista),
            ("Borrar audio", (0.78, 0.25, 0.25, 1), self.borrar_audio_desde_lista),
            ("Configuración", (0.35, 0.35, 0.35, 1), self.abrir_configuracion),
        ]:
            btn = Button(text=texto, size_hint_y=None, height=42)
            style_button(btn, color)
            btn.bind(on_press=accion)
            izquierda.add_widget(btn)

        derecha = Card(
            orientation="vertical",
            spacing=10,
            padding=14,
            bg_color=(0.12, 0.13, 0.18, 1)
        )

        derecha.add_widget(Label(
            text="¿Cómo estuvo tu día hoy?",
            font_size=26,
            size_hint_y=None,
            height=50,
            color=(1, 1, 1, 1),
            bold=True
        ))

        self.estado = Spinner(
            text="Bien",
            values=("Muy bien", "Bien", "Normal", "Cansado/a", "Triste", "Ansioso/a", "Enojado/a", "Otro"),
            size_hint_y=None,
            height=48
        )
        self.estado.background_normal = ""
        self.estado.background_color = (0.18, 0.20, 0.26, 1)
        derecha.add_widget(self.estado)

        self.titulo = TextInput(
            hint_text="Título de la nota",
            multiline=False,
            size_hint_y=None,
            height=48,
            background_color=(0.08, 0.09, 0.12, 1),
            foreground_color=(1, 1, 1, 1)
        )
        derecha.add_widget(self.titulo)

        self.texto = TextInput(
            hint_text="Escribe aquí tu nota o usa grabación/transcripción...",
            multiline=True,
            background_color=(0.08, 0.09, 0.12, 1),
            foreground_color=(1, 1, 1, 1)
        )
        derecha.add_widget(self.texto)

        grid_botones = GridLayout(cols=3, spacing=8, size_hint_y=None)
        grid_botones.bind(minimum_height=grid_botones.setter("height"))

        botones = [
            ("Nueva nota", self.nueva_nota, (0.30, 0.35, 0.85, 1)),
            ("Guardar nota", self.guardar_json, (0.18, 0.62, 0.42, 1)),
            ("Guardar dispositivo", self.guardar_en_dispositivo, (0.22, 0.44, 0.82, 1)),
            ("Guardar TXT", self.guardar_txt, (0.42, 0.42, 0.85, 1)),
            ("Guardar DOCX", self.guardar_docx, (0.52, 0.40, 0.82, 1)),
            ("Guardar PDF", self.guardar_pdf, (0.62, 0.34, 0.74, 1)),
            ("Correo", self.enviar_correo, (0.90, 0.58, 0.18, 1)),
            ("WhatsApp", self.enviar_whatsapp, (0.12, 0.70, 0.32, 1)),
            ("Grabar / Detener", self.toggle_grabacion, (0.78, 0.25, 0.25, 1)),
            ("Transcribir audio", self.transcribir_ultimo_audio_manual, (0.55, 0.20, 0.75, 1)),
            ("Transcripción en vivo", self.toggle_transcripcion_en_vivo, (0.15, 0.65, 0.65, 1)),
        ]

        self.btn_grabar = None
        self.btn_live = None

        for texto_btn, funcion, color_btn in botones:
            btn = Button(text=texto_btn, size_hint_y=None, height=50)
            style_button(btn, color_btn)
            btn.bind(on_press=funcion)
            if texto_btn == "Grabar / Detener":
                self.btn_grabar = btn
            if texto_btn == "Transcripción en vivo":
                self.btn_live = btn
            grid_botones.add_widget(btn)

        derecha.add_widget(grid_botones)

        self.status = Label(
            text="Listo",
            size_hint_y=None,
            height=35,
            color=(0.85, 0.90, 1, 1)
        )
        derecha.add_widget(self.status)

        self.add_widget(izquierda)
        self.add_widget(derecha)

        self.actualizar_historial_notas()
        self.actualizar_historial_audios()

    def _update_bg(self, *args):
        self.bg.pos = self.pos
        self.bg.size = self.size

    def slug_seguro(self, texto):
        texto = texto.strip() or "nota"
        texto = "".join(c for c in texto if c.isalnum() or c in (" ", "_", "-")).strip()
        texto = texto.replace(" ", "_")
        return texto or "nota"

    def nombre_base(self):
        titulo = self.slug_seguro(self.titulo.text)
        fecha = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        return f"{fecha}_{titulo}"

    def payload(self):
        return {
            "created_at": datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
            "feeling": self.estado.text,
            "title": self.titulo.text.strip() or "Sin título",
            "body": self.texto.text.strip()
        }

    def contenido_exportable(self):
        p = self.payload()
        return (
            f"{APP_NAME}\n"
            f"{'=' * len(APP_NAME)}\n"
            f"Fecha: {p['created_at']}\n"
            f"Estado del día: {p['feeling']}\n"
            f"Título: {p['title']}\n\n"
            f"Escrito:\n{p['body']}\n"
        )

    def validar_texto(self):
        if not self.texto.text.strip():
            self.popup("Aviso", "Escribe algo antes de guardar o enviar.")
            return False
        return True

    def popup(self, titulo, mensaje):
        caja = BoxLayout(orientation="vertical", padding=12, spacing=10)
        caja.add_widget(Label(text=mensaje))
        btn = Button(text="Cerrar", size_hint_y=None, height=45)
        style_button(btn, (0.22, 0.44, 0.82, 1))
        caja.add_widget(btn)
        pop = Popup(title=titulo, content=caja, size_hint=(0.85, 0.45))
        btn.bind(on_press=pop.dismiss)
        pop.open()

    def pedir_texto_popup(self, titulo, hint, on_accept, valor_inicial=""):
        caja = BoxLayout(orientation="vertical", padding=12, spacing=8)
        entrada = TextInput(
            text=valor_inicial,
            hint_text=hint,
            multiline=False,
            size_hint_y=None,
            height=46,
            background_color=(0.12, 0.13, 0.18, 1),
            foreground_color=(1, 1, 1, 1)
        )
        caja.add_widget(entrada)

        botones = BoxLayout(size_hint_y=None, height=45, spacing=8)
        btn_ok = Button(text="Aceptar")
        btn_cancelar = Button(text="Cancelar")
        style_button(btn_ok, (0.18, 0.62, 0.42, 1))
        style_button(btn_cancelar, (0.78, 0.25, 0.25, 1))
        botones.add_widget(btn_ok)
        botones.add_widget(btn_cancelar)
        caja.add_widget(botones)

        pop = Popup(title=titulo, content=caja, size_hint=(0.8, 0.32))

        def confirmar(_):
            on_accept(entrada.text.strip(), pop)

        btn_ok.bind(on_press=confirmar)
        btn_cancelar.bind(on_press=pop.dismiss)
        pop.open()

    # =========================
    # CONFIGURACIÓN
    # =========================
    def abrir_configuracion(self, instance):
        caja = BoxLayout(orientation="vertical", spacing=10, padding=12)

        caja.add_widget(Label(text="Idioma de transcripción", size_hint_y=None, height=28))
        idioma_spinner = Spinner(
            text=self.config_data.get("idioma_transcripcion", "es-ES"),
            values=("es-ES", "es-CL", "en-US"),
            size_hint_y=None,
            height=44
        )
        caja.add_widget(idioma_spinner)

        caja.add_widget(Label(text="Corrección automática", size_hint_y=None, height=28))
        correccion_spinner = Spinner(
            text="Sí" if self.config_data.get("correccion_automatica", True) else "No",
            values=("Sí", "No"),
            size_hint_y=None,
            height=44
        )
        caja.add_widget(correccion_spinner)

        caja.add_widget(Label(text="Punto final automático", size_hint_y=None, height=28))
        punto_spinner = Spinner(
            text="Sí" if self.config_data.get("punto_final", True) else "No",
            values=("Sí", "No"),
            size_hint_y=None,
            height=44
        )
        caja.add_widget(punto_spinner)

        caja.add_widget(Label(text="Mayúsculas automáticas", size_hint_y=None, height=28))
        mayus_spinner = Spinner(
            text="Sí" if self.config_data.get("mayusculas", True) else "No",
            values=("Sí", "No"),
            size_hint_y=None,
            height=44
        )
        caja.add_widget(mayus_spinner)

        caja.add_widget(Label(text="Carpeta de exportación", size_hint_y=None, height=28))
        carpeta_input = TextInput(
            text=self.config_data.get("carpeta_exportacion", str(EXPORTS_DIR)),
            multiline=False,
            size_hint_y=None,
            height=44,
            background_color=(0.12, 0.13, 0.18, 1),
            foreground_color=(1, 1, 1, 1)
        )
        caja.add_widget(carpeta_input)

        botones = BoxLayout(size_hint_y=None, height=45, spacing=8)
        btn_guardar = Button(text="Guardar configuración")
        btn_cancelar = Button(text="Cancelar")
        style_button(btn_guardar, (0.18, 0.62, 0.42, 1))
        style_button(btn_cancelar, (0.78, 0.25, 0.25, 1))
        botones.add_widget(btn_guardar)
        botones.add_widget(btn_cancelar)
        caja.add_widget(botones)

        pop = Popup(title="Configuración", content=caja, size_hint=(0.88, 0.86))

        def guardar_conf(_):
            self.config_data["idioma_transcripcion"] = idioma_spinner.text
            self.config_data["correccion_automatica"] = correccion_spinner.text == "Sí"
            self.config_data["punto_final"] = punto_spinner.text == "Sí"
            self.config_data["mayusculas"] = mayus_spinner.text == "Sí"
            self.config_data["carpeta_exportacion"] = carpeta_input.text.strip() or str(EXPORTS_DIR)

            Path(self.config_data["carpeta_exportacion"]).mkdir(parents=True, exist_ok=True)
            save_config(self.config_data)
            self.status.text = "Configuración guardada"
            pop.dismiss()
            self.popup("Configuración", "La configuración fue guardada correctamente.")

        btn_guardar.bind(on_press=guardar_conf)
        btn_cancelar.bind(on_press=pop.dismiss)
        pop.open()

    # =========================
    # ACENTUACIÓN / PUNTUACIÓN
    # =========================
    def corregir_acentuacion_basica(self, texto):
        texto = texto.strip()
        texto = re.sub(r"\s+", " ", texto)
        texto = re.sub(r"\s+([.,;:!?])", r"\1", texto)

        reemplazos = {
            r"\btambien\b": "también",
            r"\bmas\b": "más",
            r"\baun\b": "aún",
            r"\bmusica\b": "música",
            r"\btelefono\b": "teléfono",
            r"\bgrabacion\b": "grabación",
            r"\btranscripcion\b": "transcripción",
            r"\bdia\b": "día",
        }

        for patron, reemplazo in reemplazos.items():
            texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)

        texto = re.sub(r"\bsi\b(?=\s+(quiero|puedo|voy|me|lo|la|le|tengo|estoy|soy)\b)", "sí", texto, flags=re.IGNORECASE)
        texto = re.sub(r"\bel\b(?=\s+(dijo|quiere|puede|está|es|fue|vino)\b)", "él", texto, flags=re.IGNORECASE)
        texto = re.sub(r"\btu\b(?=\s+(eres|estás|vas|puedes|quieres|tienes)\b)", "tú", texto, flags=re.IGNORECASE)

        if self.config_data.get("punto_final", True):
            texto = self.aplicar_puntuacion_basica(texto)
        if self.config_data.get("mayusculas", True):
            texto = self.mayusculas_basicas(texto)

        return texto

    def aplicar_puntuacion_basica(self, texto):
        texto = texto.strip()

        conectores_coma = ["pero", "aunque", "sin embargo", "además", "entonces", "por eso", "por ejemplo"]
        for c in conectores_coma:
            texto = re.sub(rf"\s+({re.escape(c)})\s+", r", \1 ", texto, flags=re.IGNORECASE)

        separadores = [" luego ", " después ", " al final ", " en eso ", " de repente ", " entonces "]
        for s in separadores:
            texto = re.sub(re.escape(s), ". " + s.strip() + " ", texto, flags=re.IGNORECASE)

        preguntas_inicio = ["como", "qué", "que", "donde", "cuando", "cual", "por qué"]
        if any(texto.lower().startswith(p + " ") for p in preguntas_inicio):
            if not texto.endswith("?"):
                texto += "?"
            if not texto.startswith("¿"):
                texto = "¿" + texto[0].lower() + texto[1:]

        if texto and texto[-1] not in ".!?":
            texto += "."

        texto = re.sub(r"\.\s+\.", ".", texto)
        texto = re.sub(r",\s+,", ",", texto)
        texto = re.sub(r"\s+", " ", texto).strip()
        return texto

    def mayusculas_basicas(self, texto):
        if not texto:
            return texto
        texto = texto[0].upper() + texto[1:] if len(texto) > 1 else texto.upper()

        def repl(match):
            return match.group(1) + match.group(2).upper()

        texto = re.sub(r"([.!?]\s+)([a-záéíóúñ])", repl, texto)
        texto = re.sub(r"(¿)([a-záéíóúñ])", lambda m: m.group(1) + m.group(2).upper(), texto)
        texto = re.sub(r"(¡)([a-záéíóúñ])", lambda m: m.group(1) + m.group(2).upper(), texto)
        return texto

    def limpiar_texto_transcrito(self, texto):
        texto = texto.strip()
        texto = re.sub(r"\s+", " ", texto)
        texto = re.sub(r"\s+([,.;:!?])", r"\1", texto)

        basura = {"eh", "emm", "mmm", "ah", "este", "esto", "ok", "bueno"}
        if texto.lower() in basura:
            return ""

        if len(texto) < 3:
            return ""

        return texto

    def capitalizar_frases(self, texto):
        if not texto:
            return texto

        texto = texto[0].upper() + texto[1:] if len(texto) > 1 else texto.upper()

        def repl(match):
            return match.group(1) + match.group(2).upper()

        texto = re.sub(r"([.!?]\s+)([a-záéíóúñ])", repl, texto)
        texto = re.sub(r"(¿)([a-záéíóúñ])", lambda m: m.group(1) + m.group(2).upper(), texto)
        texto = re.sub(r"(¡)([a-záéíóúñ])", lambda m: m.group(1) + m.group(2).upper(), texto)
        return texto

    def mejorar_puntuacion_en_vivo(self, texto):
        texto = texto.strip()
        if not texto:
            return ""

        conectores = [
            "pero",
            "aunque",
            "sin embargo",
            "además",
            "entonces",
            "por eso",
            "por ejemplo",
            "la verdad",
            "después",
            "al final",
            "de repente"
        ]

        for c in conectores:
            texto = re.sub(rf"\s+({re.escape(c)})\s+", r", \1 ", texto, flags=re.IGNORECASE)

        if texto and texto[-1] not in ".!?":
            texto += "."

        texto = re.sub(r"\s+", " ", texto).strip()
        texto = self.capitalizar_frases(texto)
        return texto

    def formatear_item_transcripcion(self, texto):
        texto = self.limpiar_texto_transcrito(texto)
        if not texto:
            return ""

        if self.config_data.get("correccion_automatica", True):
            texto = self.corregir_acentuacion_basica(texto)
        else:
            texto = self.mejorar_puntuacion_en_vivo(texto)

        texto = texto.strip()
        if not texto:
            return ""

        if texto.lower() == self.live_last_text.lower():
            return ""

        self.live_last_text = texto
        return f"• {texto}"

    def actualizar_texto_con_items(self):
        base_manual = []

        contenido_actual = self.texto.text.strip()
        if contenido_actual:
            lineas = contenido_actual.splitlines()
            base_manual = [l for l in lineas if not l.strip().startswith("•")]

        bloques = []
        if base_manual:
            bloques.append("\n".join(base_manual).strip())

        if self.live_transcript_items:
            bloques.append("\n".join(self.live_transcript_items))

        self.texto.text = "\n\n".join([b for b in bloques if b.strip()])

    def agregar_item_transcrito(self, texto):
        item = self.formatear_item_transcripcion(texto)
        if not item:
            return

        self.live_transcript_items.append(item)
        self.actualizar_texto_con_items()
        self.status.text = "Texto transcrito en vivo agregado"

    # =========================
    # NOTAS
    # =========================
    def nueva_nota(self, instance):
        self.titulo.text = ""
        self.texto.text = ""
        self.estado.text = "Bien"
        self.archivo_actual = None
        self.live_transcript_items = []
        self.live_last_text = ""
        self.status.text = "Nueva nota creada"

    def guardar_json(self, instance):
        if not self.validar_texto():
            return

        data = self.payload()

        if self.archivo_actual and self.archivo_actual.exists():
            ruta = self.archivo_actual
        else:
            nombre = self.nombre_base() + ".json"
            ruta = ENTRIES_DIR / nombre
            self.archivo_actual = ruta

        ruta.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self.status.text = f"Nota guardada: {ruta.name}"
        self.actualizar_historial_notas()
        self.popup("Guardado", f"Nota guardada:\n{ruta}")

    def _export_path(self, suffix):
        export_dir = Path(self.config_data.get("carpeta_exportacion", str(EXPORTS_DIR)))
        export_dir.mkdir(parents=True, exist_ok=True)
        return export_dir / f"{self.nombre_base()}{suffix}"

    def guardar_en_dispositivo(self, instance):
        if not self.validar_texto():
            return
        ruta = self._export_path(".txt")
        ruta.write_text(self.contenido_exportable(), encoding="utf-8")
        self.status.text = f"Guardado en dispositivo: {ruta.name}"
        self.popup("Guardado", f"Archivo guardado:\n{ruta}")

    def guardar_txt(self, instance):
        if not self.validar_texto():
            return
        ruta = self._export_path(".txt")
        ruta.write_text(self.contenido_exportable(), encoding="utf-8")
        self.status.text = f"TXT guardado: {ruta.name}"
        self.popup("TXT", f"Archivo TXT guardado:\n{ruta}")

    def guardar_docx(self, instance):
        if not self.validar_texto():
            return
        if not DOCX_OK:
            self.popup("Falta librería", "Instala python-docx.")
            return

        data = self.payload()
        ruta = self._export_path(".docx")

        doc = Document()
        doc.add_heading(APP_NAME, 0)
        doc.add_paragraph(f"Fecha: {data['created_at']}")
        doc.add_paragraph(f"Estado del día: {data['feeling']}")
        doc.add_paragraph(f"Título: {data['title']}")
        doc.add_paragraph("")
        doc.add_paragraph(data["body"])
        doc.save(str(ruta))

        self.status.text = f"DOCX guardado: {ruta.name}"
        self.popup("DOCX", f"Documento guardado:\n{ruta}")

    def guardar_pdf(self, instance):
        if not self.validar_texto():
            return
        if not PDF_OK:
            self.popup("Falta librería", "Instala reportlab.")
            return

        data = self.payload()
        ruta = self._export_path(".pdf")

        c = canvas.Canvas(str(ruta), pagesize=A4)
        y = A4[1] - 50

        encabezado = [
            APP_NAME,
            f"Fecha: {data['created_at']}",
            f"Estado del día: {data['feeling']}",
            f"Título: {data['title']}",
            "",
            "Escrito:"
        ]

        for linea in encabezado:
            c.drawString(50, y, linea)
            y -= 20

        for linea in data["body"].splitlines():
            if y < 50:
                c.showPage()
                y = A4[1] - 50
            c.drawString(50, y, linea[:100])
            y -= 18

        c.save()

        self.status.text = f"PDF guardado: {ruta.name}"
        self.popup("PDF", f"Archivo PDF guardado:\n{ruta}")

    def actualizar_historial_notas(self, instance=None):
        archivos = sorted(ENTRIES_DIR.glob("*.json"), reverse=True)
        self.historial_notas = archivos

        if not archivos:
            self.lista_notas.text = "No hay notas guardadas."
            return

        lineas = []
        for i, archivo in enumerate(archivos, start=1):
            try:
                data = json.loads(archivo.read_text(encoding="utf-8"))
                titulo = data.get("title", "Sin título")
                lineas.append(f"{i}. {titulo}  [{archivo.name}]")
            except Exception:
                lineas.append(f"{i}. {archivo.name}")

        self.lista_notas.text = "\n".join(lineas)
        self.status.text = "Historial de notas actualizado"

    def abrir_nota_desde_lista(self, instance):
        if not self.historial_notas:
            self.popup("Sin notas", "No hay notas guardadas.")
            return

        def seleccionar(idx):
            archivo = self.historial_notas[idx]
            try:
                data = json.loads(archivo.read_text(encoding="utf-8"))
                self.estado.text = data.get("feeling", "Bien")
                self.titulo.text = data.get("title", "")
                self.texto.text = data.get("body", "")
                self.archivo_actual = archivo
                self.status.text = f"Nota abierta: {archivo.name}"
            except Exception as e:
                self.popup("Error", f"No se pudo abrir la nota.\n{e}")

        SelectListPopup("Selecciona una nota", self.historial_notas, seleccionar).open()

    def borrar_nota_desde_lista(self, instance):
        if not self.historial_notas:
            self.popup("Sin notas", "No hay notas guardadas.")
            return

        def seleccionar(idx):
            archivo = self.historial_notas[idx]
            try:
                archivo.unlink()
                if self.archivo_actual == archivo:
                    self.archivo_actual = None
                    self.nueva_nota(None)
                self.actualizar_historial_notas()
                self.status.text = f"Nota borrada: {archivo.name}"
                self.popup("Borrado", f"Nota eliminada:\n{archivo.name}")
            except Exception as e:
                self.popup("Error", f"No se pudo borrar la nota.\n{e}")

        SelectListPopup("Selecciona la nota a borrar", self.historial_notas, seleccionar).open()

    def editar_titulo_desde_lista(self, instance):
        if not self.historial_notas:
            self.popup("Sin notas", "No hay notas guardadas.")
            return

        def seleccionar(idx):
            archivo = self.historial_notas[idx]
            try:
                data = json.loads(archivo.read_text(encoding="utf-8"))
                titulo_actual = data.get("title", "Sin título")

                def guardar_nuevo_titulo(nuevo_titulo, pop):
                    if not nuevo_titulo:
                        self.popup("Error", "El título no puede estar vacío.")
                        return
                    data["title"] = nuevo_titulo
                    archivo.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                    if self.archivo_actual == archivo:
                        self.titulo.text = nuevo_titulo
                    self.actualizar_historial_notas()
                    self.status.text = "Título actualizado"
                    pop.dismiss()

                self.pedir_texto_popup("Editar título", "Nuevo título", guardar_nuevo_titulo, titulo_actual)
            except Exception as e:
                self.popup("Error", f"No se pudo editar el título.\n{e}")

        SelectListPopup("Selecciona la nota a editar", self.historial_notas, seleccionar).open()

    def renombrar_nota_desde_lista(self, instance):
        if not self.historial_notas:
            self.popup("Sin notas", "No hay notas guardadas.")
            return

        def seleccionar(idx):
            archivo = self.historial_notas[idx]

            def hacer_renombrado(nuevo_nombre, pop):
                if not nuevo_nombre:
                    self.popup("Error", "El nombre no puede estar vacío.")
                    return

                nuevo_slug = self.slug_seguro(nuevo_nombre)
                partes = archivo.stem.split("_", 2)
                if len(partes) >= 3:
                    nuevo_stem = f"{partes[0]}_{partes[1]}_{nuevo_slug}"
                else:
                    nuevo_stem = nuevo_slug
                nuevo_archivo = archivo.parent / f"{nuevo_stem}.json"

                try:
                    archivo.rename(nuevo_archivo)
                    if self.archivo_actual == archivo:
                        self.archivo_actual = nuevo_archivo
                    self.actualizar_historial_notas()
                    self.status.text = f"Nota renombrada: {nuevo_archivo.name}"
                    pop.dismiss()
                except Exception as e:
                    self.popup("Error", f"No se pudo renombrar la nota.\n{e}")

            self.pedir_texto_popup("Renombrar archivo", "Nuevo nombre del archivo", hacer_renombrado, archivo.stem)

        SelectListPopup("Selecciona la nota a renombrar", self.historial_notas, seleccionar).open()

    # =========================
    # AUDIOS
    # =========================
    def audio_callback(self, indata, frames, time, status):
        if status:
            print(status)
        self.audio_chunks.append(indata.copy())

    def toggle_grabacion(self, instance):
        if self.transcribiendo_en_vivo:
            self.popup("Modo activo", "Primero detén la transcripción en vivo.")
            return

        if not self.grabando:
            self.iniciar_grabacion()
        else:
            self.detener_grabacion_y_transcribir()

    def iniciar_grabacion(self):
        try:
            self.audio_chunks = []
            self.input_stream = sd.InputStream(
                samplerate=self.fs,
                channels=1,
                dtype="int16",
                callback=self.audio_callback
            )
            self.input_stream.start()
            self.grabando = True
            self.btn_grabar.text = "Detener grabación"
            style_button(self.btn_grabar, (0.95, 0.20, 0.20, 1))
            self.status.text = "Grabando audio..."
        except Exception as e:
            self.popup("Error grabando", str(e))

    def detener_grabacion_y_transcribir(self):
        try:
            if self.input_stream is not None:
                self.input_stream.stop()
                self.input_stream.close()
                self.input_stream = None

            self.grabando = False
            self.btn_grabar.text = "Grabar / Detener"
            style_button(self.btn_grabar, (0.78, 0.25, 0.25, 1))

            if not self.audio_chunks:
                self.status.text = "No se grabó audio"
                self.popup("Aviso", "No se capturó audio.")
                return

            audio_total = np.concatenate(self.audio_chunks, axis=0)
            nombre = self.nombre_base() + ".wav"
            ruta = AUDIO_DIR / nombre
            write(str(ruta), self.fs, audio_total)

            self.ultimo_audio = ruta
            self.status.text = f"Audio guardado: {nombre}"
            self.actualizar_historial_audios()
            self.transcribir_audio_automatico(ruta)

        except Exception as e:
            self.grabando = False
            self.btn_grabar.text = "Grabar / Detener"
            style_button(self.btn_grabar, (0.78, 0.25, 0.25, 1))
            self.popup("Error", str(e))

    def transcribir_ultimo_audio_manual(self, instance):
        if not self.ultimo_audio or not Path(self.ultimo_audio).exists():
            self.popup("Sin audio", "Primero debes grabar un audio.")
            return
        self.transcribir_audio_automatico(self.ultimo_audio)

    def actualizar_historial_audios(self, instance=None):
        archivos = sorted(AUDIO_DIR.glob("*.wav"), reverse=True)
        self.historial_audios = archivos

        if not archivos:
            self.lista_audios.text = "No hay audios grabados."
            return

        self.lista_audios.text = "\n".join(
            f"{i}. {archivo.name}" for i, archivo in enumerate(archivos, start=1)
        )
        self.status.text = "Lista de audios actualizada"

    def reproducir_audio_hilo(self, ruta_audio):
        try:
            self.reproduciendo = True
            data, fs = sf.read(str(ruta_audio), dtype="float32")
            sd.play(data, fs)
            sd.wait()
            Clock.schedule_once(lambda dt: self._set_status(f"Reproducción terminada: {ruta_audio.name}"), 0)
        except Exception as e:
            Clock.schedule_once(lambda dt: self.popup("Error", f"No se pudo reproducir el audio.\n{e}"), 0)
        finally:
            self.reproduciendo = False

    def reproducir_audio_desde_lista(self, instance):
        if not self.historial_audios:
            self.popup("Sin audios", "No hay audios grabados.")
            return

        def seleccionar(idx):
            archivo = self.historial_audios[idx]
            if self.reproduciendo:
                try:
                    sd.stop()
                    self.reproduciendo = False
                    self.status.text = "Reproducción detenida"
                    return
                except Exception:
                    pass

            self.status.text = f"Reproduciendo: {archivo.name}"
            hilo = threading.Thread(target=self.reproducir_audio_hilo, args=(archivo,), daemon=True)
            hilo.start()

        SelectListPopup("Selecciona el audio a reproducir", self.historial_audios, seleccionar).open()

    def borrar_audio_desde_lista(self, instance):
        if not self.historial_audios:
            self.popup("Sin audios", "No hay audios grabados.")
            return

        def seleccionar(idx):
            archivo = self.historial_audios[idx]
            try:
                archivo.unlink()
                if self.ultimo_audio == archivo:
                    self.ultimo_audio = None
                self.actualizar_historial_audios()
                self.status.text = f"Audio borrado: {archivo.name}"
                self.popup("Borrado", f"Audio eliminado:\n{archivo.name}")
            except Exception as e:
                self.popup("Error", f"No se pudo borrar el audio.\n{e}")

        SelectListPopup("Selecciona el audio a borrar", self.historial_audios, seleccionar).open()

    def transcribir_audio_automatico(self, ruta_audio):
        try:
            self.status.text = "Transcribiendo audio..."
            r = sr.Recognizer()

            with sr.AudioFile(str(ruta_audio)) as source:
                audio = r.record(source)

            texto = r.recognize_google(audio, language=self.config_data.get("idioma_transcripcion", "es-ES"))
            if self.config_data.get("correccion_automatica", True):
                texto = self.corregir_acentuacion_basica(texto)

            if self.texto.text.strip():
                self.texto.text += "\n" + texto
            else:
                self.texto.text = texto

            self.status.text = "Audio transcrito automáticamente"
            self.popup("Transcripción lista", "El audio se transcribió y se agregó a la nota.")
        except sr.UnknownValueError:
            self.status.text = "No se entendió el audio"
            self.popup("No entendí", "No se pudo entender claramente la grabación.")
        except sr.RequestError as e:
            self.status.text = "Error de transcripción"
            self.popup("Error de servicio", f"No se pudo conectar al servicio de transcripción.\n{e}")
        except Exception as e:
            self.status.text = "Error al transcribir"
            self.popup("Error", str(e))

    # =========================
    # TRANSCRIPCIÓN EN VIVO
    # =========================
    def toggle_transcripcion_en_vivo(self, instance):
        if self.grabando:
            self.popup("Modo activo", "Primero detén la grabación normal.")
            return

        if not self.transcribiendo_en_vivo:
            self.iniciar_transcripcion_en_vivo()
        else:
            self.detener_transcripcion_en_vivo()

    def live_audio_callback(self, indata, frames, time, status):
        if status:
            print(status)

        self.live_audio_buffer.append(indata.copy())

        total_frames = sum(chunk.shape[0] for chunk in self.live_audio_buffer)
        if total_frames >= self.fs * self.live_chunk_seconds:
            audio_block = np.concatenate(self.live_audio_buffer, axis=0)
            self.live_audio_buffer = []
            self.live_queue.put(audio_block)

    def live_transcription_worker(self):
        recognizer = sr.Recognizer()

        while self.transcribiendo_en_vivo or not self.live_queue.empty():
            try:
                audio_block = self.live_queue.get(timeout=0.5)
            except queue.Empty:
                continue

            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
                    temp_path = tmp.name

                write(temp_path, self.fs, audio_block)

                with sr.AudioFile(temp_path) as source:
                    audio = recognizer.record(source)

                texto = recognizer.recognize_google(
                    audio,
                    language=self.config_data.get("idioma_transcripcion", "es-ES")
                )

                texto = self.limpiar_texto_transcrito(texto)
                if not texto:
                    continue

                Clock.schedule_once(lambda dt, t=texto: self.agregar_item_transcrito(t), 0)

            except sr.UnknownValueError:
                pass
            except sr.RequestError as e:
                Clock.schedule_once(lambda dt, err=str(e): self._set_status(f"Error de red: {err}"), 0)
            except Exception as e:
                Clock.schedule_once(lambda dt, err=str(e): self._set_status(f"Error en vivo: {err}"), 0)
            finally:
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except Exception:
                        pass

    def iniciar_transcripcion_en_vivo(self):
        try:
            self.live_audio_buffer = []
            self.live_transcript_items = []
            self.live_last_text = ""

            while not self.live_queue.empty():
                try:
                    self.live_queue.get_nowait()
                except queue.Empty:
                    break

            self.input_stream = sd.InputStream(
                samplerate=self.fs,
                channels=1,
                dtype="int16",
                callback=self.live_audio_callback,
                blocksize=0
            )
            self.input_stream.start()

            self.transcribiendo_en_vivo = True
            self.btn_live.text = "Detener en vivo"
            style_button(self.btn_live, (0.95, 0.20, 0.20, 1))
            self.status.text = "Transcripción en vivo iniciada..."

            self.live_worker_thread = threading.Thread(
                target=self.live_transcription_worker,
                daemon=True
            )
            self.live_worker_thread.start()

        except Exception as e:
            self.popup("Error", f"No se pudo iniciar la transcripción en vivo.\n{e}")

    def detener_transcripcion_en_vivo(self):
        try:
            self.transcribiendo_en_vivo = False

            if self.input_stream is not None:
                self.input_stream.stop()
                self.input_stream.close()
                self.input_stream = None

            if self.live_audio_buffer:
                audio_block = np.concatenate(self.live_audio_buffer, axis=0)
                self.live_audio_buffer = []
                self.live_queue.put(audio_block)

            self.btn_live.text = "Transcripción en vivo"
            style_button(self.btn_live, (0.15, 0.65, 0.65, 1))
            self.actualizar_texto_con_items()
            self.status.text = "Transcripción en vivo detenida"

        except Exception as e:
            self.popup("Error", f"No se pudo detener la transcripción en vivo.\n{e}")

    def agregar_texto_transcrito(self, texto):
        self.agregar_item_transcrito(texto)

    def agregar_item_transcrito(self, texto):
        item = self.formatear_item_transcripcion(texto)
        if not item:
            return

        self.live_transcript_items.append(item)
        self.actualizar_texto_con_items()
        self.status.text = "Texto transcrito en vivo agregado"

    def _set_status(self, texto):
        self.status.text = texto

    # =========================
    # ENVÍO
    # =========================
    def enviar_correo(self, instance):
        if not self.validar_texto():
            return

        data = self.payload()
        asunto = quote(data["title"])
        cuerpo = quote(self.contenido_exportable())
        url = f"mailto:?subject={asunto}&body={cuerpo}"
        webbrowser.open(url)
        self.status.text = "Correo abierto"

    def enviar_whatsapp(self, instance):
        if not self.validar_texto():
            return

        mensaje = quote(self.contenido_exportable())
        url = f"https://wa.me/?text={mensaje}"
        webbrowser.open(url)
        self.status.text = "WhatsApp abierto"


class AgendaApp(App):
    def build(self):
        self.title = APP_NAME
        return AgendaLayout()


if __name__ == "__main__":
    AgendaApp().run()